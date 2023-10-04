import docx

from StringHandler import StringHandler
            
class DocumentHandler():
    def __init__(self):
        self.newDocument = None
        self.oldDocument = None
        
    def SetInputFile(self,inputFile:str):
        self.inputFile = inputFile
    def Open(self):
        self.oldDocument = docx.Document(self.inputFile)
    def SetSpellChecker(self,spellChecker:dict):
        self.spellChecker = spellChecker
    def SetSaveFullName(self,saveFullName:str):
        self.saveFullName = saveFullName
    def Save(self):
        self.oldDocument.save(self.saveFullName)
        
    def ReplaceAll(self):
        document = docx.Document()
        paragraphs = self.oldDocument.paragraphs
        for paragraph in paragraphs:
            contentList = DocumentHandler.Replace(paragraph.text,self.spellChecker)
            tempText = ''.join(contentList)
            tempParagraph = document.add_paragraph(tempText)
        self.newDocument = document
        self.oldDocument = self.newDocument
        
    
if __name__ == '__main__':
    fileFullName = "C:\\Test\\Demo\\Demo1\\inputs\\proba2.docx"
    saveFullName = "C:\\Test\\Demo\\Demo1\\Result\\output.docx"
    documentHandler = DocumentHandler()
    documentHandler.SetInputFile(fileFullName)
    documentHandler.SetSaveFullName(saveFullName)
    documentHandler.Open()
    spellChecker = {
        'Mauris' : 'Oral Sex',
        'dignissim' : 'Intercourse',
        'ultrices' : 'Group Sex'
    }
    documentHandler.SetSpellChecker(spellChecker)
    documentHandler.ReplaceAll()
    documentHandler.Save()
