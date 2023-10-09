import docx

from OfficeFiles import StringHandler

from OfficeFiles import TextHandler
            
class DocumentHandler():
    def __init__(self):
        pass
    def Load(self):
        self.document = docx.Document(self.inputFullName)
    def SetSpellChecker(self,spellChecker:dict):
        self.spellChecker = spellChecker
    def SetInputFullName(self,inputFullName : str ):
        self.inputFullName = inputFullName
    def SetSaveFullName(self,saveFullName:str):
        self.saveFullName = saveFullName
    def CopyToSaveFileName(self):
        self.SetSaveFullName(self.inputFullName)    
    def Copy(self,document):
        self.document = document 
    def AutoSave(self):
        self.document.save(self.saveFullName)
    def SearchAll(self,searchText):
        document = docx.Document()
        paragraphs = self.oldDocument.paragraphs
        resultList = list()
        for paragraph in paragraphs:
            contentList = TextHandler.TextHandler.Search(paragraph.text,searchText)
            resultList.append(contentList)
        return resultList
    
    def ReplaceAll(self):
        document = docx.Document()
        paragraphs = self.document.paragraphs
        for paragraph in paragraphs:
            contentList = TextHandler.TextHandler.Replace(paragraph.text,self.spellChecker)
            tempText = ''.join(contentList)
            tempParagraph = document.add_paragraph(tempText)
        self.Copy(document)
        self.AutoSave()
        
        
    
if __name__ == '__main__':
    fileFullName = "C:\\Test\\Demo\\Demo1\\inputs\\proba2.docx"
    saveFullName = "C:\\Test\\Demo\\Demo1\\Result\\output.docx"
    documentHandler = DocumentHandler()
    documentHandler.SetInputFile(fileFullName)
    documentHandler.SetSaveFullName(saveFullName)
    documentHandler.Load()
    spellChecker = {
        'Mauris' : 'Oral Sex',
        'dignissim' : 'Intercourse',
        'ultrices' : 'Group Sex'
    }
    documentHandler.SetSpellChecker(spellChecker)
    documentHandler.ReplaceAll()
    documentHandler.Save()
